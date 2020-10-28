'''
    Generator Application Generate Microsoft Word Document Script
    generator/generate_docx.py
    @author Teerapat Kraisrisirikul (810Teams)
'''

from datetime import datetime

from docx import Document

from clubs_and_events.settings import STORAGE_BASE_DIR, STUDENT_COMMITTEE_PRESIDENT_NAME, STUDENT_COMMITTEE_ADVISOR_NAME
from membership.models import Membership


def generate_docx(file_name, generated_file_name=None, club=None, advisor=None, objective=str(), objective_list=tuple(),
                  room=str(), schedule=str(), plan_list=tuple(), merit=str(), save=False):
    ''' Generate Microsoft Word document based on template '''
    # Init Document
    document = Document('generator/docx/{}'.format(file_name))

    # Init Generated File Name
    if generated_file_name is None:
        generated_file_name = 'generated-' + file_name

    # Init Generated Document Directories
    path = '{}/generated_docx/{}/{}'.format(STORAGE_BASE_DIR, club.id, generated_file_name)

    # Storing Data
    data = {
        'date': get_date(),
        'club_name': club.name_th.replace('ชุมนุม', ''),
        'student_committee_advisor': STUDENT_COMMITTEE_ADVISOR_NAME,
        'student_committee_president': STUDENT_COMMITTEE_PRESIDENT_NAME,
        'president': get_club_president(club),
        'advisor': advisor.name,
        'staff_list': get_staff_list(club),
        'objective': objective,
        'objective_list': objective_list,
        'room': room,
        'schedule': schedule,
        'plan_list': plan_list,
        'merit': merit,
        'member_list': get_member_list(club)
    }

    # Formatting Data - Always List
    data['objective_list'] = data['objective_list'].replace('\r', '').split('\n')
    data['plan_list'] = data['plan_list'].replace('\r', '').split('\n')

    # Formatting Data - Depends
    for i in ('schedule', 'merit'):
        if '\n' in data[i]:
            data[i] = data[i].split('\n')
        else:
            data[i] = '\t' + data[i]

    # Formatting Data - Indented String
    for i in ('advisor', 'room'):
        data[i] = '\t' + data[i]

    # Replace Paragraphs
    for paragraph in document.paragraphs:
        for inline in paragraph.runs:
            for j in data:
                if isinstance(data[j], str):
                    inline.text = inline.text.replace('{%s}' % j, data[j])
                elif isinstance(data[j], list):
                    inline.text = inline.text.replace(
                        '{%s}' % j,
                        '\n'.join(['\t{}. {}'.format(
                            i + 1, data[j][i].replace('\n', '').replace('\r', '')
                        ) for i in range(len(data[j]))])
                    )

    # Replace Paragraphs in Tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for inline in paragraph.runs:
                        for j in data:
                            if isinstance(data[j], str):
                                inline.text = inline.text.replace('{%s}' % j, data[j])

    # Document Saving
    if save:
        document.save(path)

    # Return
    return document


def get_date(date=datetime.now().date()):
    ''' Get date in a formal format '''
    months = {
        1: 'มกราคม', 2: 'กุมภาพันธ์', 3: 'มีนาคม', 4: 'เมษายน', 5: 'พฤษภาคม', 6: 'มิถุนายน',
        7: 'กรกฎาคม', 8: 'สิงหาคม', 9: 'กันยายน', 10: 'ตุลาคม', 11: 'พฤศจิกายน', 12: 'ธันวาคม',
    }

    return '{} {} {}'.format(date.day, months[date.month], date.year + 543)


def get_club_president(club):
    ''' Get club president's name '''
    membership = Membership.objects.get(community_id=club.id, position=3, status='A')

    return membership.user.name


def get_staff_list(club):
    ''' Get staff list names '''
    memberships = Membership.objects.filter(
        community_id=club.id, position__in=(1, 2, 3), status='A'
    ).order_by('-position')

    staff_list = list()

    for i in memberships:
        if i.position == 3:
            staff_list.append('{}\t\tรหัสนักศึกษา {}\t\tประธานชุมนุม'.format(i.user.name, i.user.username.replace('it', '')))
        elif i.position == 2:
            staff_list.append(
                '{}\t\tรหัสนักศึกษา {}\t\tรองประธานชุมนุม'.format(i.user.name, i.user.username.replace('it', ''))
            )
        elif i.position == 1:
            staff_list.append('{}\t\tรหัสนักศึกษา {}\t\tกรรมการ'.format(i.user.name, i.user.username.replace('it', '')))

    return staff_list


def get_member_list(club):
    ''' Get member list names '''
    memberships = Membership.objects.filter(community_id=club.id, status='A').order_by('-position')

    return [i.user.name for i in memberships]
